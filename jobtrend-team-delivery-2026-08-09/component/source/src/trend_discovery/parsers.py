from __future__ import annotations

import csv
import io
import ipaddress
import json
import mimetypes
import re
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Callable, Iterable, Mapping
from urllib.parse import urlsplit, urlunsplit

import httpx

from .io_utils import sha256_bytes, sha256_text, stable_id
from .schemas import Evidence, ExternalDocument, SourceSpec


FetchCallable = Callable[..., Any]


@dataclass(frozen=True, slots=True)
class TextSegment:
    """A traceable portion of a parsed document."""

    text: str
    char_start: int
    char_end: int
    page: int | None = None
    section: str | None = None


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    document: ExternalDocument
    segments: tuple[TextSegment, ...]


@dataclass(frozen=True, slots=True)
class _FetchedPayload:
    content: bytes
    content_type: str | None
    final_url: str


_FORMAT_BY_SUFFIX = {
    ".json": "jsonl",
    ".jsonl": "jsonl",
    ".ndjson": "jsonl",
    ".csv": "csv",
    ".txt": "txt",
    ".md": "txt",
    ".html": "html",
    ".htm": "html",
    ".pdf": "pdf",
    ".docx": "docx",
}

_FORMAT_BY_CONTENT_TYPE = {
    "application/json": "jsonl",
    "application/x-ndjson": "jsonl",
    "text/csv": "csv",
    "text/plain": "txt",
    "text/markdown": "txt",
    "text/html": "html",
    "application/xhtml+xml": "html",
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


def parse_source(
    source: SourceSpec,
    base_dir: str | Path,
    config: Mapping[str, Any],
    *,
    fetcher: FetchCallable | None = None,
    collected_at: datetime | None = None,
) -> list[ParsedDocument]:
    """Parse one manifest source into normalized, evidence-ready documents."""

    now = collected_at or datetime.now(timezone.utc)
    parsing = dict(config.get("parsing", {}))
    parser_version = str(parsing.get("parser_version", "jobtrend_parser_v1"))
    is_url = _is_url(source.input)

    if is_url:
        fetched = _fetch_url(source.input, parsing, fetcher)
        payload = fetched.content
        uri = fetched.final_url
        content_type = fetched.content_type
        name_hint = Path(urlsplit(uri).path).name or source.source_id
    else:
        path = Path(source.input).expanduser()
        if not path.is_absolute():
            path = Path(base_dir) / path
        path = path.resolve()
        payload = path.read_bytes()
        uri = None
        content_type = mimetypes.guess_type(path.name)[0]
        name_hint = path.name

    input_format = _resolve_format(
        source.input_format,
        name_hint,
        content_type,
        is_url,
        payload=payload,
    )
    if input_format == "jsonl":
        records = _read_json_records(payload)
        return [
            _structured_record(
                source,
                record,
                record_index=index,
                parser_version=parser_version,
                fetched_uri=uri,
                default_collected_at=now,
            )
            for index, record in enumerate(records, start=1)
        ]
    if input_format == "csv":
        records = _read_csv_records(payload)
        return [
            _structured_record(
                source,
                record,
                record_index=index,
                parser_version=parser_version,
                fetched_uri=uri,
                default_collected_at=now,
            )
            for index, record in enumerate(records, start=1)
        ]

    if input_format == "txt":
        title, parts, parse_status, parser_metadata = _parse_text(payload, name_hint)
    elif input_format == "html":
        title, parts, parse_status, parser_metadata = _parse_html(payload, name_hint)
    elif input_format == "pdf":
        title, parts, parse_status, parser_metadata = _parse_pdf(payload, name_hint, parsing)
    elif input_format == "docx":
        title, parts, parse_status, parser_metadata = _parse_docx(payload, name_hint, parsing)
    else:  # pragma: no cover - guarded by _resolve_format
        raise ValueError(f"unsupported input format: {input_format}")

    text, segments = _assemble_segments(parts)
    raw_hash = sha256_bytes(payload)
    external_id = str(source.metadata.get("external_id", "")).strip() or None
    document_key = external_id or source.source_id
    metadata = dict(source.metadata)
    metadata.update(
        {
            "source_id": source.source_id,
            "detected_format": input_format,
            **parser_metadata,
        }
    )
    document = ExternalDocument(
        document_id=stable_id("doc", source.source_id, document_key),
        source_type=source.source_type,
        source_name=source.source_name,
        title=title.strip() or Path(name_hint).stem or source.source_name,
        text=text,
        raw_sha256=raw_hash,
        parser_version=parser_version,
        parse_status=parse_status,
        publisher=source.publisher,
        uri=uri,
        external_id=external_id,
        published_at=source.published_at,
        collected_at=source.collected_at or now,
        license=source.license,
        metadata=metadata,
    )
    return [ParsedDocument(document=document, segments=segments)]


def build_evidence(
    parsed: ParsedDocument,
    config: Mapping[str, Any],
) -> list[Evidence]:
    """Split parsed segments while retaining exact offsets into document.text."""

    document = parsed.document
    if document.parse_status != "parsed" or not document.text:
        return []

    parsing = dict(config.get("parsing", {}))
    chunk_chars = int(parsing.get("chunk_chars", 1200))
    overlap = int(parsing.get("chunk_overlap", 120))
    if chunk_chars <= 0:
        raise ValueError("parsing.chunk_chars must be positive")
    if overlap < 0 or overlap >= chunk_chars:
        raise ValueError("parsing.chunk_overlap must be in [0, chunk_chars)")

    segments = parsed.segments or (
        TextSegment(document.text, 0, len(document.text)),
    )
    evidence: list[Evidence] = []
    for segment in segments:
        cursor = segment.char_start
        while cursor < segment.char_end:
            proposed_end = min(cursor + chunk_chars, segment.char_end)
            end = _natural_boundary(document.text, cursor, proposed_end, segment.char_end)
            trimmed_start = cursor
            while trimmed_start < end and document.text[trimmed_start].isspace():
                trimmed_start += 1
            trimmed_end = end
            while trimmed_end > trimmed_start and document.text[trimmed_end - 1].isspace():
                trimmed_end -= 1
            if trimmed_end > trimmed_start:
                chunk = document.text[trimmed_start:trimmed_end]
                chunk_hash = sha256_text(chunk)
                evidence.append(
                    Evidence(
                        evidence_id=stable_id(
                            "ev",
                            document.document_id,
                            segment.page or "",
                            segment.section or "",
                            trimmed_start,
                            trimmed_end,
                            chunk_hash,
                        ),
                        document_id=document.document_id,
                        source_type=document.source_type,
                        text=chunk,
                        text_sha256=chunk_hash,
                        uri=document.uri,
                        page=segment.page,
                        section=segment.section,
                        char_start=trimmed_start,
                        char_end=trimmed_end,
                    )
                )
            if end >= segment.char_end:
                break
            cursor = max(cursor + 1, end - overlap)
    return evidence


def _structured_record(
    source: SourceSpec,
    record: Mapping[str, Any],
    *,
    record_index: int,
    parser_version: str,
    fetched_uri: str | None,
    default_collected_at: datetime,
) -> ParsedDocument:
    lookup = {str(key).strip().casefold(): value for key, value in record.items()}
    title = _text_value(
        _first(
            lookup,
            "job_title",
            "title",
            "position_name",
            "position",
            "job_name",
            "name",
            "岗位名称",
            "职位名称",
            "职位",
        )
    )
    company = _text_value(
        _first(lookup, "company_name", "company", "employer", "公司名称", "公司", "企业名称")
    )
    industry = _text_value(_first(lookup, "industry", "industry_name", "行业", "所属行业"))
    region = _text_value(
        _first(lookup, "location", "region", "city", "work_location", "工作地点", "地区", "城市")
    )
    external_id = _text_value(
        _first(lookup, "job_id", "position_id", "external_id", "id", "uuid", "职位id", "岗位id")
    )
    record_uri = _text_value(
        _first(lookup, "url", "uri", "job_url", "source_url", "link", "detail_url", "职位链接")
    )
    published_value = _first(
        lookup,
        "publish_date",
        "published_at",
        "published_date",
        "publication_date",
        "发布日期",
        "发布时间",
    )
    collected_value = _first(
        lookup,
        "scrape_time",
        "collected_at",
        "collection_time",
        "crawled_at",
        "crawl_time",
        "采集时间",
        "抓取时间",
    )

    body = _text_value(
        _first(
            lookup,
            "jd_text",
            "job_description",
            "description",
            "content",
            "jd",
            "job_detail",
            "职位描述",
            "职位详情",
            "岗位描述",
        )
    )
    responsibilities = _text_value(
        _first(lookup, "responsibilities", "responsibility", "duties", "岗位职责", "工作职责")
    )
    requirements = _text_value(
        _first(lookup, "requirements", "requirement", "qualifications", "任职要求", "岗位要求")
    )
    skills = _text_value(
        _first(lookup, "skills_norm", "skills_raw", "skills", "skill", "技能", "技能要求")
    )
    if not body:
        sections: list[str] = []
        if responsibilities:
            sections.append(f"岗位职责：\n{responsibilities}")
        if requirements:
            sections.append(f"岗位要求：\n{requirements}")
        if skills:
            sections.append(f"技能：\n{skills}")
        body = "\n\n".join(sections)
    if not title:
        title = f"{company + ' - ' if company else ''}{source.source_name}记录{record_index}"

    canonical_record = json.dumps(record, ensure_ascii=False, sort_keys=True, default=str, separators=(",", ":"))
    raw_hash = sha256_text(canonical_record)
    document_key = external_id or raw_hash
    published_at = _datetime_value(published_value) or source.published_at
    collected_at = _datetime_value(collected_value) or source.collected_at or default_collected_at

    metadata = dict(source.metadata)
    metadata.update({"source_id": source.source_id, "record_index": record_index, "detected_format": "record"})
    structured = {
        "responsibilities": responsibilities,
        "requirements": requirements,
        "skills": skills,
    }
    metadata["structured_fields"] = {key: value for key, value in structured.items() if value}

    normalized_body = _normalize_text(body)
    text, segments = _assemble_segments([(normalized_body, None, title)])
    document = ExternalDocument(
        document_id=stable_id("doc", source.source_id, document_key),
        source_type=source.source_type,
        source_name=source.source_name,
        title=title,
        text=text,
        raw_sha256=raw_hash,
        parser_version=parser_version,
        parse_status="parsed",
        publisher=_text_value(_first(lookup, "publisher", "发布单位")) or source.publisher,
        uri=record_uri or fetched_uri,
        external_id=external_id,
        company=company,
        industry=industry,
        region=region,
        published_at=published_at,
        collected_at=collected_at,
        license=source.license,
        metadata=metadata,
    )
    return ParsedDocument(document=document, segments=segments)


def _read_json_records(payload: bytes) -> list[dict[str, Any]]:
    text = _decode_text(payload)
    stripped = text.lstrip()
    if stripped.startswith("["):
        value = json.loads(text)
        if not isinstance(value, list) or not all(isinstance(row, dict) for row in value):
            raise ValueError("JSON input must contain an array of objects")
        return list(value)

    records: list[dict[str, Any]] = []
    for line_number, line in enumerate(text.splitlines(), start=1):
        if not line.strip():
            continue
        try:
            value = json.loads(line)
        except json.JSONDecodeError as exc:
            raise ValueError(f"JSONL line {line_number} is invalid: {exc.msg}") from exc
        if not isinstance(value, dict):
            raise ValueError(f"JSONL line {line_number} must be an object")
        records.append(value)
    return records


def _read_csv_records(payload: bytes) -> list[dict[str, str]]:
    reader = csv.DictReader(io.StringIO(_decode_text(payload)))
    if not reader.fieldnames:
        raise ValueError("CSV header is missing")
    return [{str(key): str(value or "") for key, value in row.items()} for row in reader]


def _parse_text(
    payload: bytes,
    name_hint: str,
) -> tuple[str, list[tuple[str, int | None, str | None]], str, dict[str, Any]]:
    text = _normalize_text(_decode_text(payload))
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    title = lines[0].lstrip("# ") if lines else Path(name_hint).stem
    parts = _split_heading_sections(text)
    return title, parts, "parsed", {"parser_backend": "text"}


def _parse_html(
    payload: bytes,
    name_hint: str,
) -> tuple[str, list[tuple[str, int | None, str | None]], str, dict[str, Any]]:
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:  # pragma: no cover - dependency is declared by the package
        raise RuntimeError("beautifulsoup4 is required to parse HTML") from exc

    soup = BeautifulSoup(payload, "html.parser")
    for node in soup(["script", "style", "noscript", "template", "nav", "footer", "aside"]):
        node.decompose()
    title = ""
    if soup.title and soup.title.get_text(strip=True):
        title = soup.title.get_text(" ", strip=True)
    heading = soup.find(["h1", "h2"])
    if not title and heading:
        title = heading.get_text(" ", strip=True)
    title = title or Path(name_hint).stem

    container = None
    for selector in (
        "article",
        "main",
        '[role="main"]',
        ".TRS_Editor",
        ".article-content",
        ".article_content",
        ".content-detail",
        ".detail-content",
        "#UCAP-CONTENT",
        "#content",
    ):
        container = soup.select_one(selector)
        if container is not None:
            break
    container = container or soup.body or soup
    parts: list[tuple[str, int | None, str | None]] = []
    active_heading: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        value = _normalize_text("\n".join(buffer))
        if value:
            parts.append((value, None, active_heading))
        buffer = []

    for element in container.find_all(["h1", "h2", "h3", "h4", "p", "li", "pre"], recursive=True):
        value = element.get_text(" ", strip=True)
        if not value:
            continue
        if element.name in {"h1", "h2", "h3", "h4"}:
            flush()
            active_heading = value
        else:
            buffer.append(value)
    flush()
    fallback = _normalize_text(container.get_text("\n", strip=True))
    extracted_chars = sum(len(item[0]) for item in parts)
    # Government pages sometimes put the article in nested div/span nodes
    # without semantic paragraphs.  Prefer the lossless container text when
    # the paragraph-only pass captured almost nothing.  Header/nav/footer were
    # removed above so this fallback does not silently become site chrome.
    if not parts or (len(fallback) >= 200 and len(fallback) > extracted_chars * 2):
        parts = [(fallback, None, title)] if fallback else []
    return title, parts, "parsed", {"parser_backend": "beautifulsoup4"}


def _parse_pdf(
    payload: bytes,
    name_hint: str,
    parsing: Mapping[str, Any],
) -> tuple[str, list[tuple[str, int | None, str | None]], str, dict[str, Any]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared by the package
        raise RuntimeError("pypdf is required to parse PDF files") from exc

    reader = PdfReader(io.BytesIO(payload))
    title = ""
    if reader.metadata:
        title = str(getattr(reader.metadata, "title", "") or "").strip()
    title = title or Path(name_hint).stem
    parts: list[tuple[str, int | None, str | None]] = []
    non_whitespace_chars = 0
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            page_text = _normalize_text(page.extract_text() or "")
        except Exception:
            page_text = ""
        non_whitespace_chars += len(re.sub(r"\s+", "", page_text))
        if page_text:
            parts.append((page_text, page_number, f"第{page_number}页"))

    page_count = len(reader.pages)
    minimum_per_page = int(parsing.get("min_pdf_text_chars_per_page", 40))
    needs_ocr = page_count == 0 or non_whitespace_chars < minimum_per_page * max(1, page_count)
    docling_text = _docling_text(payload, ".pdf", parsing)
    # Docling is used as a layout/text fallback only for non-scanned PDFs. A PDF
    # that fails the text-coverage gate remains excluded from automatic scoring.
    if not parts and docling_text and not needs_ocr:
        parts = [(docling_text, None, title)]
    metadata = {
        "parser_backend": "pypdf" + ("+docling" if docling_text else ""),
        "page_count": page_count,
        "text_char_count": non_whitespace_chars,
    }
    return title, parts, "needs_ocr" if needs_ocr else "parsed", metadata


def _parse_docx(
    payload: bytes,
    name_hint: str,
    parsing: Mapping[str, Any],
) -> tuple[str, list[tuple[str, int | None, str | None]], str, dict[str, Any]]:
    docling_text = _docling_text(payload, ".docx", parsing)
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is declared by the package
        if docling_text:
            title = Path(name_hint).stem
            return title, [(docling_text, None, title)], "parsed", {"parser_backend": "docling"}
        raise RuntimeError("python-docx is required to parse DOCX files") from exc

    document = Document(io.BytesIO(payload))
    title = str(document.core_properties.title or "").strip()
    parts: list[tuple[str, int | None, str | None]] = []
    active_heading: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        value = _normalize_text("\n".join(buffer))
        if value:
            parts.append((value, None, active_heading))
        buffer = []

    for paragraph in document.paragraphs:
        value = _normalize_text(paragraph.text)
        if not value:
            continue
        style_name = str(getattr(paragraph.style, "name", "") or "").casefold()
        if style_name.startswith("heading") or style_name.startswith("标题"):
            flush()
            active_heading = value
            title = title or value
        else:
            buffer.append(value)
    for table in document.tables:
        rows = [" | ".join(_normalize_text(cell.text) for cell in row.cells) for row in table.rows]
        table_text = _normalize_text("\n".join(rows))
        if table_text:
            buffer.append(table_text)
    flush()
    title = title or Path(name_hint).stem
    if not parts and docling_text:
        parts = [(docling_text, None, title)]
    backend = "python-docx" + ("+docling" if docling_text else "")
    return title, parts, "parsed", {"parser_backend": backend}


def _docling_text(payload: bytes, suffix: str, parsing: Mapping[str, Any]) -> str | None:
    if parsing.get("use_docling", True) is False:
        return None
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        return None

    temporary_path: Path | None = None
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as handle:
            handle.write(payload)
            temporary_path = Path(handle.name)
        result = DocumentConverter().convert(str(temporary_path))
        value = result.document.export_to_markdown()
        return _normalize_text(str(value)) or None
    except Exception:
        return None
    finally:
        if temporary_path is not None:
            temporary_path.unlink(missing_ok=True)


def _assemble_segments(
    parts: Iterable[tuple[str, int | None, str | None]],
) -> tuple[str, tuple[TextSegment, ...]]:
    output: list[str] = []
    segments: list[TextSegment] = []
    length = 0
    for raw_text, page, section in parts:
        value = _normalize_text(raw_text)
        if not value:
            continue
        if output:
            output.append("\n\n")
            length += 2
        start = length
        output.append(value)
        length += len(value)
        segments.append(
            TextSegment(
                text=value,
                char_start=start,
                char_end=length,
                page=page,
                section=section,
            )
        )
    return "".join(output), tuple(segments)


def _split_heading_sections(text: str) -> list[tuple[str, int | None, str | None]]:
    lines = text.splitlines()
    parts: list[tuple[str, int | None, str | None]] = []
    active: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        nonlocal buffer
        value = _normalize_text("\n".join(buffer))
        if value:
            parts.append((value, None, active))
        buffer = []

    for line in lines:
        stripped = line.strip()
        is_markdown_heading = bool(re.match(r"^#{1,6}\s+\S", stripped))
        is_short_cn_heading = bool(re.match(r"^.{1,30}[：:]$", stripped))
        if is_markdown_heading or is_short_cn_heading:
            flush()
            active = stripped.lstrip("# ").rstrip("：:").strip()
        else:
            buffer.append(line)
    flush()
    if not parts and text:
        parts.append((text, None, None))
    return parts


def _natural_boundary(text: str, start: int, proposed_end: int, hard_end: int) -> int:
    if proposed_end >= hard_end:
        return hard_end
    minimum = start + max(1, (proposed_end - start) * 3 // 5)
    candidates = [
        text.rfind(marker, minimum, proposed_end)
        for marker in ("\n", "。", "！", "？", ". ", "; ", "；")
    ]
    boundary = max(candidates, default=-1)
    return boundary + 1 if boundary >= minimum else proposed_end


def _fetch_url(url: str, parsing: Mapping[str, Any], fetcher: FetchCallable | None) -> _FetchedPayload:
    _validate_public_url(url)
    timeout = float(parsing.get("request_timeout_seconds", 60))
    user_agent = str(parsing.get("user_agent", "jobtrend-research-bot/0.1"))
    headers = {"User-Agent": user_agent, "Accept": "*/*"}
    if fetcher is None:
        response = httpx.get(url, timeout=timeout, headers=headers, follow_redirects=True)
    else:
        try:
            response = fetcher(url, timeout=timeout, headers=headers)
        except TypeError:
            response = fetcher(url)
    result = _coerce_fetch_result(response, url)
    _validate_public_url(result.final_url)
    return result


def _coerce_fetch_result(value: Any, original_url: str) -> _FetchedPayload:
    if isinstance(value, bytes):
        return _FetchedPayload(value, None, original_url)
    if isinstance(value, str):
        return _FetchedPayload(value.encode("utf-8"), "text/plain", original_url)
    if isinstance(value, tuple) and len(value) == 2:
        body, content_type = value
        content = body if isinstance(body, bytes) else str(body).encode("utf-8")
        return _FetchedPayload(content, str(content_type) if content_type else None, original_url)
    if isinstance(value, Mapping):
        status = int(value.get("status_code", 200))
        if status >= 400:
            raise RuntimeError(f"remote source returned HTTP {status}")
        body = value.get("content", value.get("body", value.get("text", b"")))
        content = body if isinstance(body, bytes) else str(body).encode("utf-8")
        headers = value.get("headers", {})
        content_type = headers.get("content-type") if isinstance(headers, Mapping) else None
        final_url = str(value.get("url", original_url))
        return _FetchedPayload(content, str(content_type) if content_type else None, final_url)

    if hasattr(value, "raise_for_status"):
        value.raise_for_status()
    status_code = int(getattr(value, "status_code", 200))
    if status_code >= 400:
        raise RuntimeError(f"remote source returned HTTP {status_code}")
    body = getattr(value, "content", None)
    if body is None:
        body = str(getattr(value, "text", "")).encode("utf-8")
    elif not isinstance(body, bytes):
        body = bytes(body)
    headers = getattr(value, "headers", {})
    content_type = headers.get("content-type") if isinstance(headers, Mapping) else None
    final_url = str(getattr(value, "url", original_url))
    return _FetchedPayload(body, str(content_type) if content_type else None, final_url)


def _validate_public_url(value: str) -> None:
    parsed = urlsplit(value)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise ValueError("URL sources must use http or https")
    if parsed.username or parsed.password:
        raise ValueError("URL sources must not contain credentials")
    try:
        address = ipaddress.ip_address(parsed.hostname.strip("[]"))
    except ValueError:
        return
    if not address.is_global:
        raise ValueError("URL sources must target a public address")


def _resolve_format(
    declared: str,
    name_hint: str,
    content_type: str | None,
    is_url: bool,
    *,
    payload: bytes | None = None,
) -> str:
    if declared not in {"auto", "url"}:
        return declared

    # Some authoritative download endpoints return an incorrect generic MIME
    # type (for example OSTA's PDF endpoint may report an Excel MIME type), and
    # their URLs do not necessarily have a useful filename suffix.  Binary
    # signatures are therefore a stronger signal than response headers here.
    detected = _format_from_payload(payload)
    if detected:
        return detected

    clean_content_type = str(content_type or "").split(";", 1)[0].strip().casefold()
    if clean_content_type in _FORMAT_BY_CONTENT_TYPE:
        return _FORMAT_BY_CONTENT_TYPE[clean_content_type]
    suffix = Path(name_hint).suffix.casefold()
    if suffix in _FORMAT_BY_SUFFIX:
        return _FORMAT_BY_SUFFIX[suffix]
    return "html" if is_url else "txt"


def _format_from_payload(payload: bytes | None) -> str | None:
    """Return a conservative format inference from content signatures.

    Only unambiguous PDF and DOCX signatures override a server-provided MIME
    type.  HTML is recognised when a text payload begins with a normal HTML
    declaration/tag; this keeps opaque binary downloads out of the HTML parser.
    """

    if not payload:
        return None
    if payload.startswith(b"%PDF-"):
        return "pdf"
    if payload.startswith(b"PK\x03\x04") and _is_docx_zip(payload):
        return "docx"

    prefix = payload[:1024].lstrip(b"\xef\xbb\xbf\x00\t\r\n ").lower()
    if prefix.startswith((b"<!doctype html", b"<html", b"<head", b"<body", b"<main", b"<article")):
        return "html"
    return None


def _is_docx_zip(payload: bytes) -> bool:
    """Detect DOCX packages without treating every ZIP download as a DOCX."""

    try:
        with zipfile.ZipFile(io.BytesIO(payload)) as archive:
            members = set(archive.namelist())
    except (OSError, zipfile.BadZipFile):
        return False
    return "[Content_Types].xml" in members and "word/document.xml" in members


def _is_url(value: str) -> bool:
    return urlsplit(value).scheme.casefold() in {"http", "https"}


def _first(values: Mapping[str, Any], *aliases: str) -> Any:
    for alias in aliases:
        value = values.get(alias.casefold())
        if not _is_empty(value):
            return value
    return None


def _is_empty(value: Any) -> bool:
    if value is None:
        return True
    if isinstance(value, str) and value.strip().casefold() in {"", "null", "none", "nan", "nat"}:
        return True
    return False


def _text_value(value: Any) -> str | None:
    if _is_empty(value):
        return None
    if isinstance(value, str):
        return _normalize_text(value)
    if isinstance(value, Mapping):
        return _normalize_text(json.dumps(value, ensure_ascii=False, sort_keys=True, default=str))
    if isinstance(value, Iterable) and not isinstance(value, (bytes, bytearray)):
        return _normalize_text("\n".join(str(item) for item in value if not _is_empty(item)))
    return _normalize_text(str(value))


def _datetime_value(value: Any) -> datetime | None:
    if _is_empty(value):
        return None
    if isinstance(value, datetime):
        return value
    if isinstance(value, date):
        return datetime.combine(value, datetime.min.time())
    if isinstance(value, (int, float)):
        seconds = float(value)
        if seconds > 10_000_000_000:
            seconds /= 1000
        try:
            return datetime.fromtimestamp(seconds, tz=timezone.utc)
        except (OverflowError, OSError, ValueError):
            return None
    text = str(value).strip()
    if not text:
        return None
    normalized = text.replace("Z", "+00:00")
    try:
        return datetime.fromisoformat(normalized)
    except ValueError:
        pass
    for pattern in ("%Y/%m/%d %H:%M:%S", "%Y/%m/%d", "%Y年%m月%d日"):
        try:
            return datetime.strptime(text, pattern)
        except ValueError:
            continue
    return None


def _decode_text(payload: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return payload.decode(encoding)
        except UnicodeDecodeError:
            continue
    return payload.decode("utf-8", errors="replace")


def _normalize_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    value = re.sub(r"[ \t]+\n", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def canonical_public_url(value: str) -> str:
    """Return a fragment-free URL suitable for source metadata and stable matching."""

    parsed = urlsplit(value)
    return urlunsplit((parsed.scheme.casefold(), parsed.netloc.casefold(), parsed.path, parsed.query, ""))


__all__ = [
    "FetchCallable",
    "ParsedDocument",
    "TextSegment",
    "build_evidence",
    "canonical_public_url",
    "parse_source",
]
