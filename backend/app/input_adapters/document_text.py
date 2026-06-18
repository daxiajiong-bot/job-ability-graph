"""Document text extraction adapters.

The matching and resume parsers work on plain text. This module isolates file
format handling so PDF/Word support can evolve without coupling it to the
algorithm pipeline.
"""

from __future__ import annotations

import dataclasses
from pathlib import Path
from typing import Dict, Iterable, Optional


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


class DocumentExtractionError(ValueError):
    """Raised when a document cannot be converted to text."""


@dataclasses.dataclass
class DocumentText:
    text: str
    metadata: Dict[str, object]


def extract_text_from_path(path: str | Path, filename: Optional[str] = None) -> DocumentText:
    file_path = Path(path)
    if not file_path.exists():
        raise DocumentExtractionError(f"file does not exist: {file_path}")
    return extract_text_from_bytes(file_path.read_bytes(), filename=filename or file_path.name)


def extract_text_from_bytes(content: bytes, filename: str) -> DocumentText:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentExtractionError(
            f"unsupported document type '{extension or '<none>'}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    if not content:
        raise DocumentExtractionError("document content is empty")

    if extension in {".txt", ".md"}:
        text = _decode_text(content)
        method = "plain_text"
        page_count = None
    elif extension == ".pdf":
        text, page_count = _extract_pdf_text(content)
        method = "pypdf"
    elif extension == ".docx":
        text, page_count = _extract_docx_text(content)
        method = "python-docx"
    else:
        raise DocumentExtractionError(f"unsupported document type: {extension}")

    text = normalize_extracted_text(text)
    if not text:
        raise DocumentExtractionError("no extractable text found in document")
    return DocumentText(
        text=text,
        metadata={
            "filename": filename,
            "extension": extension,
            "extractor": method,
            "page_count": page_count,
            "char_count": len(text),
            "supported_extensions": sorted(SUPPORTED_EXTENSIONS),
        },
    )


def normalize_extracted_text(text: str) -> str:
    lines = [line.strip() for line in (text or "").splitlines()]
    compacted = []
    previous_blank = False
    for line in lines:
        if not line:
            if not previous_blank:
                compacted.append("")
            previous_blank = True
            continue
        compacted.append(line)
        previous_blank = False
    return "\n".join(compacted).strip()


def _decode_text(content: bytes) -> str:
    encodings = ("utf-8", "utf-8-sig", "gb18030", "latin-1")
    for encoding in encodings:
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentExtractionError("failed to decode text document")


def _extract_pdf_text(content: bytes) -> tuple[str, Optional[int]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise DocumentExtractionError("pypdf is required for PDF text extraction") from exc

    import io

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError("failed to read PDF document") from exc
    page_texts = []
    for page in reader.pages:
        page_texts.append(page.extract_text() or "")
    return "\n".join(page_texts), len(reader.pages)


def _extract_docx_text(content: bytes) -> tuple[str, Optional[int]]:
    try:
        import docx
    except ImportError as exc:
        raise DocumentExtractionError("python-docx is required for DOCX text extraction") from exc

    import io

    try:
        document = docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise DocumentExtractionError("failed to read DOCX document") from exc
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    parts.extend(_table_text(document.tables))
    return "\n".join(parts), None


def _table_text(tables: Iterable[object]) -> list[str]:
    rows = []
    for table in tables:
        for row in getattr(table, "rows", []):
            cells = [cell.text.strip() for cell in getattr(row, "cells", []) if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
    return rows
